#!/usr/bin/env python3
"""
山西财经大学本科学年论文自动格式化工具
依据：附件1《山西财经大学本科学年论文内容与格式规范》

用法：python format_paper.py <论文.docx> [年级]
示例：python format_paper.py 论文主体.docx 2024
"""

import sys
import os
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# 格式参数
# ============================================================
PAGE_TOP = Cm(3.0)
PAGE_BOTTOM = Cm(2.5)
PAGE_LEFT = Cm(2.5)
PAGE_RIGHT = Cm(2.0)

BODY_FONT = '宋体'
BODY_SIZE = Pt(12)
BODY_INDENT = Cm(0.74)
BODY_LINE_SPACING = 1.25

L1_FONT = '黑体'
L1_SIZE = Pt(16)
L1_ALIGN = WD_ALIGN_PARAGRAPH.CENTER
L1_BEFORE = Pt(12)
L1_AFTER = Pt(12)

L2_FONT = '宋体'
L2_SIZE = Pt(14)
L2_ALIGN = WD_ALIGN_PARAGRAPH.LEFT

L3_FONT = '宋体'
L3_SIZE = Pt(12)
L3_ALIGN = WD_ALIGN_PARAGRAPH.LEFT

HEADER_FONT = '宋体'
HEADER_SIZE = Pt(9)


def detect(text):
    """
    检测文本首行标题级别。
    返回: (level, heading_text) 或 (None, None)

    标题类型:
      一级(L1): "1  引  言"、"一、xxx"、摘要/Abstract/目录/参考文献/致谢/附录
      二级(L2): "1.1  xxx"、"（一）xxx"
      三级(L3): "1.1.1  xxx"
    """
    if not text:
        return None, None

    first_line = text.strip().split('\n')[0].strip()

    # --- L1: "1  引  言" (Arabic number + 2+ spaces + content) ---
    if re.match(r'^\d+\s{2,}\S', first_line) and len(first_line) < 40:
        return 1, first_line

    # --- L1: "一、xxx" (Chinese numeral + 、) ---
    if re.match(r'^[一二三四五六七八九十]+、', first_line) and len(first_line) < 30:
        return 1, first_line

    # --- L1: 特殊节名称 ---
    if first_line in ['摘要', '摘 要', 'Abstract', 'ABSTRACT',
                       '目录', '目 录', '参考文献', '参考 文献',
                       '致谢', '致 谢', '附录', '附 录',
                       '导论', '导 论', '结语', '结 语']:
        return 1, first_line

    # --- L3: "1.1.1  xxx" ---
    if re.match(r'^\d+\.\d+\.\d+\s', first_line) and len(first_line) < 50:
        return 3, first_line

    # --- L2: "1.1  xxx" ---
    if re.match(r'^\d+\.\d+\s', first_line) and len(first_line) < 40:
        return 2, first_line

    # --- L2: "（一）xxx" ---
    if re.match(r'^（[一二三四五六七八九十]+）', first_line) and len(first_line) < 30:
        return 2, first_line

    return None, None


def is_ref(text):
    """检测是否为参考文献条目: [n] xxx"""
    return bool(re.match(r'^\[\d+\]', text.strip()))


def setup_pages(doc, grade):
    """页面设置 + 页眉"""
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = PAGE_TOP
        section.bottom_margin = PAGE_BOTTOM
        section.left_margin = PAGE_LEFT
        section.right_margin = PAGE_RIGHT
        section.different_first_page_header_footer = True

        header = section.header
        header.is_linked_to_previous = False
        for p in header.paragraphs:
            p.clear()
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.paragraph_format.space_before = Pt(0)
        hp.paragraph_format.space_after = Pt(0)
        run = hp.add_run(f'山西财经大学{grade}级本科生学年论文')
        run.font.size = HEADER_SIZE
        run.font.name = HEADER_FONT
        run._element.rPr.rFonts.set(qn('w:eastAsia'), HEADER_FONT)


def fmt_runs(para, font_name, font_size, bold, alignment):
    """格式化段落中所有 run"""
    pf = para.paragraph_format
    pf.line_spacing = BODY_LINE_SPACING
    para.alignment = alignment
    for run in para.runs:
        run.font.size = font_size
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
        run.font.bold = bold


def split_body(doc):
    """
    将论文正文长段落按标题边界拆分为独立段落。
    论文正文常是一个巨段，标题和正文混在一起，必须先拆分。
    """
    heading_boundary = re.compile(r'\n(?=\d+(?:\.\d+)*\s{2,}\S)')

    modified = True
    passes = 0
    while modified and passes < 10:
        modified = False
        passes += 1

        for i in range(len(doc.paragraphs) - 1, -1, -1):
            para = doc.paragraphs[i]
            text = para.text

            if len(text) < 200:
                continue

            # 检查是否包含多个标题
            test_text = '\n' + text
            matches = heading_boundary.findall(test_text)
            if len(matches) < 2:
                continue

            # 按标题边界拆分
            sections = re.split(heading_boundary, test_text)
            sections = [s.strip() for s in sections if s.strip()]
            if len(sections) <= 1:
                continue

            parent = para._element.getparent()
            insert_idx = list(parent).index(para._element)

            new_elems = []
            for section_text in sections:
                p_elem = OxmlElement('w:p')
                r_elem = OxmlElement('w:r')
                t_elem = OxmlElement('w:t')
                t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                t_elem.text = section_text
                r_elem.append(t_elem)
                p_elem.append(r_elem)
                new_elems.append(p_elem)

            for elem in reversed(new_elems):
                parent.insert(insert_idx + 1, elem)

            parent.remove(para._element)
            modified = True
            break

    return doc


def format_all(doc):
    """对所有段落应用格式"""
    counts = {'l1': 0, 'l2': 0, 'l3': 0, 'body': 0, 'ref': 0}

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

        level, _ = detect(text)

        if level == 1:
            pf.first_line_indent = Cm(0)
            pf.space_before = L1_BEFORE
            pf.space_after = L1_AFTER
            fmt_runs(para, L1_FONT, L1_SIZE, True, L1_ALIGN)
            counts['l1'] += 1

        elif level == 2:
            pf.first_line_indent = Cm(0)
            fmt_runs(para, L2_FONT, L2_SIZE, True, L2_ALIGN)
            counts['l2'] += 1

        elif level == 3:
            pf.first_line_indent = Cm(0)
            fmt_runs(para, L3_FONT, L3_SIZE, True, L3_ALIGN)
            counts['l3'] += 1

        elif is_ref(text):
            pf.first_line_indent = Cm(0)
            fmt_runs(para, BODY_FONT, BODY_SIZE, False, WD_ALIGN_PARAGRAPH.LEFT)
            counts['ref'] += 1

        else:
            pf.first_line_indent = BODY_INDENT
            fmt_runs(para, BODY_FONT, BODY_SIZE, False, WD_ALIGN_PARAGRAPH.LEFT)
            counts['body'] += 1

    return counts


def format_paper(input_path, grade="20XX"):
    """
    主入口：格式化学年论文

    Args:
        input_path: .docx 论文路径
        grade: 年级 (如 "2024")
    Returns:
        output_path: 输出文件路径
    """
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"File not found: {input_path}")

    dir_name = os.path.dirname(input_path)
    base_name = os.path.splitext(os.path.basename(input_path))[0]
    output_path = os.path.join(dir_name, f"{base_name}_formatted.docx")

    doc = Document(input_path)

    # Step 1: Split giant body paragraphs
    doc = split_body(doc)

    # Step 2: Page setup + headers
    setup_pages(doc, grade)

    # Step 3: Apply formatting to all paragraphs
    counts = format_all(doc)

    # Step 4: Save
    doc.save(output_path)

    # Report
    print(f"Output: {output_path}")
    print(f"Page: A4, margins 3.0/2.5/2.5/2.0 cm")
    print(f"Body: SimSun 12pt, 1.25x line spacing, first-line indent 2 chars")
    print(f"L1: HeiTi 16pt bold centered, L2: SimSun 14pt bold left")
    print(f"Header: SimSun 9pt - Shanxi Univ. of Finance and Economics")
    print(f"Counts: L1={counts['l1']}, L2={counts['l2']}, L3={counts['l3']}, "
          f"body={counts['body']}, ref={counts['ref']}")

    return output_path


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python format_paper.py <paper.docx> [grade]")
        print("Example: python format_paper.py mypaper.docx 2024")
        sys.exit(1)

    input_path = sys.argv[1]
    grade = sys.argv[2] if len(sys.argv) > 2 else "20XX"

    try:
        format_paper(input_path, grade)
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)
