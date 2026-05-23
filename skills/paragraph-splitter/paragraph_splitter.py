#!/usr/bin/env python3
"""
段落拆分器 - 解决.docx中换行符过滤导致的段落粘连问题

用法：
  python paragraph_splitter.py <文件.docx> --detect    # 仅检测
  python paragraph_splitter.py <文件.docx> --split      # 检测并拆分
  python paragraph_splitter.py <文件.docx> --report     # 拆分后报告
"""

import sys
import os
import re
import argparse
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# 拆分规则：在段落文本中的 \n 处匹配这些模式
# ============================================================
SPLIT_PATTERNS = [
    # L1: "1  引  言"  (数字 + 2+空格 + 汉字)
    r'\n(?=\d+\s{2,}[一-龥])',
    # L2: "1.1  研究背景"  (数字.数字 + 空格 + 汉字)
    r'\n(?=\d+\.\d+\s+[一-龥A-Z])',
    # L3: "1.1.1  方法"
    r'\n(?=\d+\.\d+\.\d+\s+[一-龥A-Z])',
    # 中文序号: "一、xxx"
    r'\n(?=[一二三四五六七八九十]+、)',
    # 中文括号: "（一）xxx"
    r'\n(?=（[一二三四五六七八九十]+）)',
    # 特殊节: 摘要/Abstract/目录/参考文献/致谢/附录/导论/结语
    r'\n(?=(?:摘要|Abstract|ABSTRACT|目录|参考文献|致谢|附录|导论|结语)(?:\s|$))',
    # 参考文献条目: "[1] xxx"
    r'\n(?=\[\d+\]\s)',
]

# 拆分模式（合并所有规则）
SPLIT_RE = re.compile('|'.join(f'({p})' for p in SPLIT_PATTERNS))


def detect_stuck_paragraphs(doc, min_length=200):
    """检测粘连段落，返回 [(index, length, sections_count)]"""
    stuck = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if len(text) >= min_length:
            # 计算内部包含多少个可拆分点
            test = '\n' + text
            boundaries = SPLIT_RE.findall(test)
            if len(boundaries) >= 2:
                stuck.append((i, len(text), len(boundaries)))
    return stuck


def split_stuck_paragraph(doc, para_index):
    """拆分指定索引的粘连段落，返回新段落数"""
    para = doc.paragraphs[para_index]
    text = para.text

    if len(text) < 50:
        return 0

    test = '\n' + text
    boundaries = SPLIT_RE.findall(test)
    if len(boundaries) < 2:
        return 0

    # 拆分
    sections = re.split(SPLIT_RE, test)
    # re.split with captures returns interleaved groups; filter empty/None
    cleaned = []
    for s in sections:
        if s and s.strip():
            cleaned.append(s.strip())

    if len(cleaned) <= 1:
        return 0

    parent = para._element.getparent()
    insert_idx = list(parent).index(para._element)

    new_elems = []
    for section_text in cleaned:
        p_elem = OxmlElement('w:p')
        r_elem = OxmlElement('w:r')
        t_elem = OxmlElement('w:t')
        t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t_elem.text = section_text
        r_elem.append(t_elem)
        p_elem.append(r_elem)
        new_elems.append(p_elem)

    for elem in reversed(new_elems):
        parent.insert(insert_idx + 1, elem)

    parent.remove(para._element)
    return len(cleaned)


def split_all(doc, min_length=200):
    """拆分文档中所有粘连段落"""
    results = []
    # Iterate in reverse to handle indices correctly
    stuck = detect_stuck_paragraphs(doc, min_length)
    # Process from end to start to preserve indices
    stuck.sort(key=lambda x: x[0], reverse=True)

    for idx, orig_len, boundary_count in stuck:
        n = split_stuck_paragraph(doc, idx)
        if n > 0:
            results.append((idx, orig_len, n))

    return results


def generate_report(doc):
    """生成段落统计报告"""
    total = len(doc.paragraphs)
    non_empty = sum(1 for p in doc.paragraphs if p.text.strip())
    long = sum(1 for p in doc.paragraphs if len(p.text) > 500)
    giant = sum(1 for p in doc.paragraphs if len(p.text) > 5000)

    l1 = sum(1 for p in doc.paragraphs
             if re.match(r'^\d+\s{2,}\S', p.text.strip().split('\n')[0].strip()))
    l2 = sum(1 for p in doc.paragraphs
             if re.match(r'^\d+\.\d+\s', p.text.strip().split('\n')[0].strip()))

    return {
        'total': total,
        'non_empty': non_empty,
        'long_paragraphs': long,
        'giant_paragraphs': giant,
        'l1_headings': l1,
        'l2_headings': l2,
    }


def main():
    parser = argparse.ArgumentParser(description='段落拆分器 - 解决.docx段落粘连')
    parser.add_argument('input', help='输入.docx文件路径')
    parser.add_argument('--detect', action='store_true', help='仅检测粘连段落')
    parser.add_argument('--split', action='store_true', help='执行拆分')
    parser.add_argument('--report', action='store_true', help='拆分后统计报告')
    parser.add_argument('--min-length', type=int, default=200,
                        help='最小粘连检测字符数 (默认: 200)')
    parser.add_argument('--output', '-o', help='输出文件路径 (默认: 原名_split.docx)')
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"Error: file not found: {args.input}")
        sys.exit(1)

    doc = Document(args.input)

    if args.detect:
        stuck = detect_stuck_paragraphs(doc, args.min_length)
        print(f"Found {len(stuck)} stuck paragraphs:")
        for idx, length, boundaries in stuck:
            print(f"  P{idx}: {length} chars, ~{boundaries} possible sections")
        return

    if args.split:
        results = split_all(doc, args.min_length)
        print(f"Split {len(results)} paragraphs:")
        total_sections = 0
        for idx, orig_len, n in results:
            print(f"  P{idx}: {orig_len} chars -> {n} sections")
            total_sections += n

        output = args.output or args.input.replace('.docx', '_split.docx')
        doc.save(output)
        print(f"\nSaved: {output}")
        print(f"Total new sections created: {total_sections}")
        return

    if args.report:
        report = generate_report(doc)
        print(f"Total paragraphs: {report['total']}")
        print(f"Non-empty: {report['non_empty']}")
        print(f"Long (>500 chars): {report['long_paragraphs']}")
        print(f"Giant (>5000 chars): {report['giant_paragraphs']}")
        print(f"L1 headings detected: {report['l1_headings']}")
        print(f"L2 headings detected: {report['l2_headings']}")
        return

    # Default: detect + report
    stuck = detect_stuck_paragraphs(doc, args.min_length)
    report = generate_report(doc)
    print(f"Document: {args.input}")
    print(f"Paragraphs: {report['total']} total, {report['non_empty']} non-empty")
    print(f"Stuck paragraphs: {len(stuck)}")
    if stuck:
        for idx, length, boundaries in stuck:
            print(f"  P{idx}: {length} chars, ~{boundaries} sections")
    print(f"Run with --split to fix, --report for details")


if __name__ == '__main__':
    main()
